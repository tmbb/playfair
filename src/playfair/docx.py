import docx
import os
import math

def new_styled_document():
    directory = os.path.dirname(__file__)
    path = os.path.join(directory, 'docx/styled-doc.docx')
    return docx.Document(path)


def _format_as_plaintext(formatter, value):
    doc = docx.Document()
    paragraph = doc.add_paragraph()
    formatter.insert(paragraph, value)
    plaintext = paragraph.text
    return plaintext


def _p_value_relation(paragraph, sign, p_value):
    paragraph.add_run("p").italic = True
    paragraph.add_run(" {} ".format(sign))
    paragraph.add_run(p_value)


def _pretty_format_float(value, func):
    if value == float('inf'):
        return '+∞'
    elif value == (- float('inf')):
        return '-∞'
    elif math.isnan(value):
        return 'NaN'
    else:
        return func(value)

def _format_float_as_integer(value):
    return _pretty_format_float(value, lambda f: str(int(f)))

def _round_float(f, nr_of_places):
    return ("{:.%if}" % nr_of_places).format(f)

def _format_rounded_float(value, nr_of_places):
    return _pretty_format_float(value, lambda f: _round_float(f, nr_of_places))

class DocxFormatter:

    def as_plaintext(self, value):
        return _format_as_plaintext(self, value)

class IntegerFormatter(DocxFormatter):

    def __init__(self):
        pass

    def insert(self, paragraph, value):
        text = _format_float_as_integer(value)
        paragraph.add_run(text)
        return paragraph

class DefaultFormatter(DocxFormatter):

    def __init__(self):
        pass

    def insert(self, paragraph, value):
        text = str(value)
        paragraph.add_run(text)
        return paragraph

class RoundedFormatter(DocxFormatter):

    def __init__(self, nr_of_places):
        self.nr_of_places = nr_of_places

    def insert(self, paragraph, value):
        text = _format_rounded_float(value, self.nr_of_places)
        paragraph.add_run(text)
        return paragraph


class TruncatedPValueFormatter(DocxFormatter):

    def __init__(self, steps_less_than=[0.0001, 0.001, 0.01, 0.05], steps_greater_than=[], nr_of_places=3):
        self.steps_less_than = sorted(steps_less_than)
        self.steps_greater_than = sorted(steps_greater_than, reverse=True)
        self.nr_of_places = nr_of_places

    def insert(self, paragraph, value):
        for step in self.steps_less_than:
            if value < step:
                p_value = _format_rounded_float(step, self.nr_of_places)
                _p_value_relation(paragraph, "<", p_value)
                return paragraph

        for step in self.steps_greater_than:
            if value > step:
                p_value = _format_rounded_float(step, self.nr_of_places)
                _p_value_relation(paragraph, ">", p_value)
                return paragraph

        p_value = _format_rounded_float(value, self.nr_of_places)
        _p_value_relation(paragraph, "=", p_value)
        return paragraph


class TruncatedFormatter(DocxFormatter):

    def __init__(self, steps_less_than=[], steps_greater_than=[], nr_of_places=3):
        self.steps_less_than = sorted(steps_less_than)
        self.steps_greater_than = sorted(steps_greater_than, reverse=True)
        self.nr_of_places = nr_of_places

    def insert(self, paragraph, value):
        for step in self.steps_less_than:
            if value < step:
                rounded_value = _format_rounded_float(step, self.nr_of_places)
                text = "< {}".format(rounded_value)
                paragraph.add_run(text)
                return paragraph

        for step in self.steps_greater_than:
            if value > step:
                rounded_value = _format_rounded_float(step, self.nr_of_places)
                text = "> {}".format(rounded_value)
                paragraph.add_run(text)
                return paragraph

        p_value = _format_rounded_float(value, self.nr_of_places)
        paragraph.add_run(p_value)
        return paragraph


def add_table_from_dataframe(document, dataframe, formatters=dict(), strip_index=True, caption=None):
    table = document.add_table(rows=(dataframe.shape[0] + 1), cols=dataframe.shape[1])
    # Insert columns into table
    # Column names are simple strings, so this doesn't need much customization
    for i, column in enumerate(dataframe.columns):
        paragraph = table.cell(0, i).paragraphs[0]
        paragraph.add_run(column).bold = True

    # Dataframe values may be arbitrarily processed before being written
    for i, column in enumerate(dataframe) :
        for row in range(dataframe.shape[0]):
            # See if we have a special inserter/formatter
            formatter = formatters.get(column, DefaultFormatter())
            # Extract the value from the table
            value = dataframe[column][row]
            # Get the paragraphs from the table cell
            # (there may be more than one!)
            paragraphs = table.cell(row + 1, i).paragraphs
            formatter.insert(paragraph, value)